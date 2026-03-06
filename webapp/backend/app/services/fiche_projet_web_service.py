"""
fiche_projet_web_service.py
Génère la fiche projet Word (.docx) pour la version web (PostgreSQL).
Utilise python-docx — codes couleurs identiques au modèle DSI :
  #C00000 rouge      → en-têtes sections majeures
  #F2DDDC rose pâle  → zones de saisie / valeurs
  #D8D8D8 gris clair → libellés
  #BFBFBF gris moyen → sous-sections
  #F2F2F2 gris léger → alternance lignes
"""
import os
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

RED   = 'C00000'
PINK  = 'F2DDDC'
GRAY1 = 'D8D8D8'
GRAY2 = 'BFBFBF'
GRAY4 = 'F2F2F2'
WHITE = 'FFFFFF'
BLACK = '000000'


# ─────────────────────────────────────────────────────────────────────────────
# Helpers XML bas-niveau
# ─────────────────────────────────────────────────────────────────────────────
def _rgb(h):
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _shade(tc, fill):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = tc._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def _margins(tc, v=55, h=110):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = tc._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for side, val in [('top', v), ('bottom', v), ('left', h), ('right', h)]:
        n = OxmlElement(f'w:{side}')
        n.set(qn('w:w'), str(val))
        n.set(qn('w:type'), 'dxa')
        m.append(n)
    tcPr.append(m)


def _no_border(tc):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = tc._tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        n = OxmlElement(f'w:{side}')
        n.set(qn('w:val'), 'nil')
        b.append(n)
    tcPr.append(b)


def _row_height(row, twips):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight')
    h.set(qn('w:val'), str(twips))
    h.set(qn('w:hRule'), 'atLeast')
    trPr.append(h)


def _cell_width(tc, twips):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = tc._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcW'))
    if existing is not None:
        existing.set(qn('w:w'), str(twips))
        existing.set(qn('w:type'), 'dxa')
    else:
        w = OxmlElement('w:tcW')
        w.set(qn('w:w'), str(twips))
        w.set(qn('w:type'), 'dxa')
        tcPr.insert(0, w)


def _tbl_width(tbl, twips):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    existing = tblPr.find(qn('w:tblW'))
    if existing is not None:
        existing.set(qn('w:w'), str(twips))
        existing.set(qn('w:type'), 'dxa')
    else:
        w = OxmlElement('w:tblW')
        w.set(qn('w:w'), str(twips))
        w.set(qn('w:type'), 'dxa')
        tblPr.append(w)


def _spacing(p_obj, before=25, after=25):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p_obj._p.get_or_add_pPr()
    existing = pPr.find(qn('w:spacing'))
    if existing is not None:
        existing.set(qn('w:before'), str(before))
        existing.set(qn('w:after'), str(after))
    else:
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'), str(before))
        sp.set(qn('w:after'), str(after))
        rPr = pPr.find(qn('w:rPr'))
        if rPr is not None:
            rPr.addprevious(sp)
        else:
            pPr.append(sp)


# ─────────────────────────────────────────────────────────────────────────────
# Moteur de construction de lignes
# ─────────────────────────────────────────────────────────────────────────────
def _build_row(tbl, cells_cfg, default_height=None):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL

    row = tbl.add_row()
    widths = tbl._col_widths
    raw_cells = list(row.cells)
    col_idx = 0

    for cd in cells_cfg:
        if col_idx >= len(raw_cells):
            break
        tc = raw_cells[col_idx]
        span = cd.get('span', 1)

        if span > 1:
            end = min(col_idx + span, len(raw_cells))
            for j in range(col_idx + 1, end):
                tc.merge(raw_cells[j])

        _cell_width(tc, sum(widths[col_idx:col_idx + span]))
        _shade(tc, cd.get('fill', WHITE))
        if cd.get('no_border'):
            _no_border(tc)
        _margins(tc)
        tc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p_obj = tc.paragraphs[0]
        p_obj.clear()
        p_obj.alignment = cd.get('align', WD_ALIGN_PARAGRAPH.LEFT)
        _spacing(p_obj)

        run = p_obj.add_run(str(cd.get('text', '') or ''))
        run.font.name = 'Arial'
        run.font.size = Pt(cd.get('size', 8))
        run.font.bold = cd.get('bold', False)
        run.font.italic = cd.get('italic', False)
        run.font.color.rgb = RGBColor(*_rgb(cd.get('color', BLACK)))

        col_idx += max(span, 1)

    h = default_height or next((cd.get('height') for cd in cells_cfg if cd.get('height')), None)
    if h:
        _row_height(row, h)
    return row


def _make_table(doc, col_widths):
    tbl = doc.add_table(rows=0, cols=len(col_widths))
    tbl.style = 'Table Grid'
    _tbl_width(tbl, sum(col_widths))
    tbl._col_widths = col_widths
    return tbl


# ─────────────────────────────────────────────────────────────────────────────
# Raccourcis cellules
# ─────────────────────────────────────────────────────────────────────────────
def _sec(text, span=4):
    return dict(text=text, fill=RED, bold=True, size=9, color=WHITE, span=span, no_border=True)

def _sub(text, span=4, fill=GRAY2):
    return dict(text=text, fill=fill, bold=True, size=8, color=BLACK, span=span, no_border=True)

def _lbl(text):
    return dict(text=text, fill=GRAY1, bold=True, size=8, color=BLACK)

def _val(text, fill=PINK):
    return dict(text=str(text) if text is not None else '', fill=fill, size=8, color=BLACK)

def _chdr(text):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    return dict(text=text, fill=RED, bold=True, size=8, color=WHITE,
                no_border=True, align=WD_ALIGN_PARAGRAPH.CENTER)

def _ctr(text, fill=PINK):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    return dict(text=str(text) if text is not None else '', fill=fill, size=8,
                align=WD_ALIGN_PARAGRAPH.CENTER)

def _empty(span=4, h=600):
    return dict(text='', fill=WHITE, size=8, span=span, height=h)

def _fmt_eur(v):
    try:
        return f"{float(v or 0):,.2f} €"
    except Exception:
        return '0,00 €'

def _fmt_pct(v):
    try:
        return f"{int(v or 0)} %"
    except Exception:
        return '0 %'

def _fmt_date(d):
    if not d:
        return ''
    try:
        return datetime.fromisoformat(str(d)[:10]).strftime('%d/%m/%Y')
    except Exception:
        return str(d)[:10]


# ─────────────────────────────────────────────────────────────────────────────
# GÉNÉRATION WORD
# ─────────────────────────────────────────────────────────────────────────────
def generer_fiche_projet(data: dict, output_path: str) -> str:
    """
    Génère le fichier Word de la fiche projet.
    data        : dict avec toutes les infos du projet
    output_path : chemin de sortie du .docx
    """
    from docx import Document
    from docx.shared import Pt, Twips, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import json as _json

    doc = Document()

    sec = doc.sections[0]
    sec.page_width   = Twips(12240)
    sec.page_height  = Twips(15840)
    sec.left_margin  = sec.right_margin  = Twips(720)
    sec.top_margin   = sec.bottom_margin = Twips(720)
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(8)

    p = data
    today = datetime.now().strftime('%d/%m/%Y')

    CW4  = [2700, 2700, 2700, 2700]
    CW4b = [2000, 3400, 2000, 3400]
    CW_T = [3600, 2200, 2600, 2400]
    CW_E = [2160, 2700, 3240, 2700]

    def br():
        doc.add_paragraph()

    # ── BANDEAU TITRE ─────────────────────────────────────────────────────────
    t = _make_table(doc, [10800])
    _build_row(t, [dict(text='FICHE PROJET — Budget Manager Pro V5',
                        fill=RED, bold=True, size=12, color=WHITE,
                        align=WD_ALIGN_PARAGRAPH.CENTER, no_border=True)], default_height=450)
    _build_row(t, [dict(text=f"Date : {today}   |   Code : {p.get('code','')}   |   Statut : {p.get('statut','')}",
                        fill=RED, size=8, color=PINK,
                        align=WD_ALIGN_PARAGRAPH.CENTER, no_border=True)])
    br()

    # ── 1. IDENTIFICATION ──────────────────────────────────────────────────────
    t = _make_table(doc, CW4b)
    _build_row(t, [_sec('1. IDENTIFICATION')])
    _build_row(t, [_sub('Informations générales')])
    _build_row(t, [_lbl('Code projet'),   _val(p.get('code')),
                   _lbl('Intitulé'),      _val(p.get('nom'))])
    _build_row(t, [_lbl('Type'),          _val(p.get('type_projet')),
                   _lbl('Priorité'),      _val(p.get('priorite'))])
    _build_row(t, [_lbl('Phase'),         _val(p.get('phase')),
                   _lbl('Avancement'),    _val(_fmt_pct(p.get('avancement')))])
    _build_row(t, [_lbl('Date début'),    _val(_fmt_date(p.get('date_debut'))),
                   _lbl('Fin prévue'),    _val(_fmt_date(p.get('date_fin')))])
    _build_row(t, [_lbl('Fin réelle'),    _val(_fmt_date(p.get('date_fin_reelle'))),
                   _lbl('Service'),       _val(p.get('service', ''))])

    _build_row(t, [_sub('Description détaillée du projet')])
    _build_row(t, [dict(text=p.get('description', ''), fill=PINK, size=8, span=4, height=850)])

    _build_row(t, [_sub('Acteurs du projet')])
    _build_row(t, [_chdr('Rôle'), _chdr('Nom Prénom'), _chdr('Fonction / Service'), _chdr('Email / Tél')])
    for role, nom in [('Chef de projet DSI', p.get('chef_projet', '')),
                      ('Responsable métier',  p.get('responsable', '')),
                      ('Équipe projet',       p.get('equipe', '')),
                      ('Prestataires',        p.get('prestataires', ''))]:
        _build_row(t, [_lbl(role), _val(nom, WHITE), _val('', WHITE), _val('', PINK)])
    br()

    # ── 2. OPPORTUNITÉ ─────────────────────────────────────────────────────────
    t = _make_table(doc, CW4)
    _build_row(t, [_sec('2. OPPORTUNITE')])
    _build_row(t, [_sub('Objectifs "métier" opérationnels du projet')])
    _build_row(t, [_sub('Objectifs identifiés', span=2, fill=GRAY1),
                   _sub('Description détaillée', span=2, fill=GRAY1)])
    _build_row(t, [dict(text=p.get('objectifs', ''), fill=PINK, size=8, span=2, height=550),
                   _empty(span=2, h=550)])

    _build_row(t, [_sub('Principaux risques identifiés à NE PAS faire le projet')])
    _build_row(t, [_sub('Risques identifiés', span=2, fill=GRAY1),
                   _sub('Description détaillée', span=2, fill=GRAY1)])
    _build_row(t, [dict(text=p.get('risques', ''), fill=PINK, size=8, span=2, height=500),
                   _empty(span=2, h=500)])

    _build_row(t, [_sub('Gains qualitatifs et bénéfices attendus')])
    _build_row(t, [dict(text=p.get('gains', ''), fill=PINK, size=8, span=4, height=500)])

    _build_row(t, [_sub("Enjeux stratégiques de l'établissement")])
    _build_row(t, [dict(text=p.get('enjeux', ''), fill=PINK, size=8, span=4, height=400)])

    _build_row(t, [_sub("Liens avec d'autres projets internes et/ou externes")])
    _build_row(t, [_sub('Projets SI internes', span=2, fill=GRAY1),
                   _sub('Projets externes / régionaux / nationaux', span=2, fill=GRAY1)])
    _build_row(t, [_empty(span=2, h=500), _empty(span=2, h=500)])
    br()

    # ── 3. BUDGET & FINANCEMENT ────────────────────────────────────────────────
    t = _make_table(doc, CW4)
    _build_row(t, [_sec('3. BUDGET & FINANCEMENT')])
    _build_row(t, [_sub('Synthèse budgétaire')])
    _build_row(t, [_lbl('Budget prévisionnel'), _val(_fmt_eur(p.get('budget_previsionnel'))),
                   _lbl('Budget voté'),          _val(_fmt_eur(p.get('budget_vote')))])
    _build_row(t, [_lbl('Budget consommé (BC)'), _val(_fmt_eur(p.get('budget_consomme'))),
                   _lbl('Ligne budgétaire V5'),  _val(p.get('ligne_budg', ''))])
    _build_row(t, [_lbl('Bons de commande liés'),
                   dict(text=p.get('bcs', 'Aucun BC lié'), fill=PINK, size=8, span=3)])

    _build_row(t, [_sub('Coûts et charges par phase')])
    _build_row(t, [_chdr('Catégorie'), _chdr('Définition projet'),
                   _chdr('Mise en œuvre'), _chdr('Total €')])
    couts = p.get('couts_detail', {}) or {}
    total_def = total_meo = 0.0
    for cat in ['MOE interne', 'MOA interne', 'Licences / Logiciels',
                'Materiels / Serveurs', 'Sous-traitance', 'Autres']:
        cd = couts.get(cat, {})
        if not cd:
            for k in couts:
                if (k.lower().replace('é', 'e').replace('è', 'e').replace('â', 'a') ==
                        cat.lower().replace('é', 'e').replace('è', 'e').replace('â', 'a')):
                    cd = couts[k]
                    break
        def_ = cd.get('definition', 0) or 0
        meo  = cd.get('mise_en_oeuvre', 0) or 0
        tot  = cd.get('total', def_ + meo) or 0
        total_def += def_
        total_meo += meo
        _build_row(t, [dict(text=cat, fill=GRAY4, size=8),
                       _ctr(f"{def_:,.2f}" if def_ else '', WHITE),
                       _ctr(f"{meo:,.2f}"  if meo  else '', WHITE),
                       _ctr(f"{tot:,.2f}"  if tot  else '', PINK)])
    grand_total = total_def + total_meo
    _build_row(t, [dict(text='TOTAL', fill=GRAY2, bold=True, size=8),
                   _ctr(f"{total_def:,.2f}" if total_def else '', PINK),
                   _ctr(f"{total_meo:,.2f}" if total_meo else '', PINK),
                   dict(text=f"{grand_total:,.2f}" if grand_total else '', fill=PINK, bold=True, size=8)])

    _build_row(t, [_sub('Modalités de financement (internes / externes)')])
    fin_text = p.get('financement', '')
    _build_row(t, [dict(text=fin_text or '', fill=PINK if fin_text else WHITE,
                        size=8, span=4, height=500)])
    br()

    # ── 4. PLANNING & TÂCHES ──────────────────────────────────────────────────
    taches = p.get('taches', []) or []
    t = _make_table(doc, CW_T)
    _build_row(t, [_sec('4. PLANNING & TACHES')])
    _build_row(t, [_chdr('Tâche / Livrable'), _chdr('Statut'),
                   _chdr('Échéance'), _chdr('Charge (h)')])
    if taches:
        for i, tk in enumerate(taches):
            f = GRAY4 if i % 2 else WHITE
            _build_row(t, [dict(text=tk.get('titre', ''), fill=f, size=8),
                           _ctr(tk.get('statut', ''), f),
                           _ctr(tk.get('echeance', ''), PINK),
                           _ctr(tk.get('heures', ''), PINK)])
    else:
        for _ in range(4):
            _build_row(t, [_empty(span=4, h=450)])
    br()

    # ── 5. ÉQUIPE & CONTACTS ──────────────────────────────────────────────────
    contacts = p.get('contacts', []) or []
    t = _make_table(doc, CW_E)
    _build_row(t, [_sec('5. EQUIPE & CONTACTS')])
    _build_row(t, [_chdr('Rôle'), _chdr('Nom Prénom'),
                   _chdr('Fonction / Service'), _chdr('Email / Tél')])

    acteurs = [('Chef de projet DSI', p.get('chef_projet', ''), 'DSI', ''),
               ('Responsable métier',  p.get('responsable', ''), '', '')]
    for c in contacts:
        acteurs.append((c.get('role', ''), c.get('nom', ''),
                        c.get('fonction', ''), c.get('email', '')))
    if not contacts and p.get('equipe'):
        acteurs.append(('Équipe projet', p.get('equipe', ''), '', ''))

    for role, nom, fn, email in acteurs:
        _build_row(t, [_lbl(role), _val(nom, WHITE), _val(fn, WHITE), _val(email, PINK)])
    br()

    # ── 6. SOLUTIONS & CONTRAINTES ────────────────────────────────────────────
    t = _make_table(doc, CW4)
    _build_row(t, [_sec('6. SOLUTIONS & CONTRAINTES')])
    _build_row(t, [_sub('Ressources MOE disponibles — Solutions envisagées')])
    _build_row(t, [dict(text=p.get('solutions', ''), fill=PINK, size=8, span=4,
                        height=max(600, 200 + 80 * len((p.get('solutions') or '').split('\n'))))])
    _build_row(t, [_sub('Contraintes techniques / réglementaires / RGPD')])
    _build_row(t, [dict(text=p.get('contraintes', ''), fill=PINK, size=8, span=4,
                        height=max(500, 200 + 80 * len((p.get('contraintes') or '').split('\n'))))])
    br()

    # ── 6b. LES 6 CONTRAINTES PROJET ─────────────────────────────────────────
    ctr_data = {}
    ctr_raw = p.get('contraintes_6axes')
    if ctr_raw:
        try:
            ctr_data = _json.loads(ctr_raw) if isinstance(ctr_raw, str) else (ctr_raw or {})
        except Exception:
            pass

    t = _make_table(doc, CW4)
    _build_row(t, [_sec('6b. LES 6 CONTRAINTES PROJET')])
    _build_row(t, [_sub('Les 6 contraintes sont interconnectées — modifier l\'une impacte les autres.', span=4)])

    AXES_6 = [
        ('portee_desc',      'Portée',     'Périmètre, livrables, fonctionnalités incluses'),
        ('couts_desc',       'Coûts',      'Budget global, salaires, équipements, licences'),
        ('delais_desc',      'Délais',     'Calendrier, jalons, phases, dates clés'),
        ('ressources_desc',  'Ressources', 'Équipes, compétences, équipements, logiciels'),
        ('qualite_desc',     'Qualité',    "Critères d'acceptation, niveaux de service"),
        ('risques_proj_desc','Risques',    'Événements imprévus pouvant impacter le projet'),
    ]
    for i in range(0, len(AXES_6), 2):
        row_cells = []
        for j in range(2):
            if i + j < len(AXES_6):
                attr, label, hint = AXES_6[i + j]
                val = ctr_data.get(attr, '') or ''
                txt = f"{label}\n{val if val else hint}"
                row_cells.append(dict(text=txt, fill=PINK if val else GRAY4,
                                      size=8, span=2,
                                      height=max(400, 150 + 60 * len(val.split('\n')))))
            else:
                row_cells.append(_empty(span=2, h=400))
        _build_row(t, row_cells)
    br()

    # ── 6c. TRIANGLE D'OR ─────────────────────────────────────────────────────
    tri_data = {}
    tri_raw = p.get('triangle_tensions')
    if tri_raw:
        try:
            tri_data = _json.loads(tri_raw) if isinstance(tri_raw, str) else (tri_raw or {})
        except Exception:
            pass

    t = _make_table(doc, CW4)
    _build_row(t, [_sec("6c. TRIANGLE D'OR — PORTÉE / COÛTS / DÉLAIS")])
    _build_row(t, [_sub('Niveau de tension sur chaque axe (1 = sous contrôle, 5 = très tendu)', span=4)])

    NIVEAUX = {1: 'Faible', 2: 'Modéré', 3: 'Modéré', 4: 'Élevé', 5: 'Critique'}
    COULEURS_TRI = {1: '27AE60', 2: '27AE60', 3: 'F39C12', 4: 'E67E22', 5: 'C0392B'}

    axes_tri = [
        ('tension_portee', 'Portée',  'Risque de dérive du périmètre (scope creep)'),
        ('tension_couts',  'Coûts',   'Pression sur le budget disponible'),
        ('tension_delais', 'Délais',  'Pression sur le calendrier'),
    ]
    _build_row(t, [_chdr('Axe'), _chdr('Tension (1-5)'), _chdr('Niveau'), _chdr('Interprétation')])
    for attr, label, desc in axes_tri:
        v = int(tri_data.get(attr, 3))
        col = COULEURS_TRI.get(v, 'F39C12')
        niv = NIVEAUX.get(v, 'Modéré')
        barres = '[' + '|' * v + ' ' * (5 - v) + ']'
        _build_row(t, [_lbl(label),
                       dict(text=f"{v}/5  {barres}", fill=col,
                            bold=True, size=9, color='FFFFFF'),
                       _val(niv, PINK),
                       dict(text=desc, fill=WHITE, size=8)])

    total_tri = sum(int(tri_data.get(a, 3)) for a, _, _ in axes_tri)
    if total_tri >= 13:
        alerte_tri = 'ALERTE : Triangle très tendu — arbitrage urgent nécessaire !'
        fill_tri = 'C0392B'
    elif total_tri >= 10:
        alerte_tri = 'Attention : tensions élevées sur certains axes — surveiller de près.'
        fill_tri = 'E67E22'
    else:
        alerte_tri = 'Triangle équilibré — les trois axes sont sous contrôle.'
        fill_tri = '27AE60'
    _build_row(t, [dict(text=alerte_tri, fill=fill_tri,
                        bold=True, size=9, span=4, color='FFFFFF')])

    arbitrage = p.get('arbitrage', '')
    if arbitrage:
        _build_row(t, [_sub("Stratégie d'arbitrage", span=4)])
        _build_row(t, [dict(text=arbitrage, fill=PINK, size=8, span=4,
                            height=max(400, 150 + 60 * len(arbitrage.split('\n'))))])
    br()

    # ── 6d. REGISTRE DES RISQUES ──────────────────────────────────────────────
    risques_list = []
    rsk_raw = p.get('registre_risques')
    if rsk_raw:
        try:
            risques_list = _json.loads(rsk_raw) if isinstance(rsk_raw, str) else (rsk_raw or [])
        except Exception:
            pass

    t = _make_table(doc, [3800, 1400, 800, 700, 700, 3000, 1200])
    _build_row(t, [_sec('6d. REGISTRE DES RISQUES', span=7)])
    _build_row(t, [_chdr('Description'), _chdr('Catégorie'), _chdr('Proba'),
                   _chdr('Impact'), _chdr('Crit.'), _chdr('Action corrective'), _chdr('Statut')])

    if risques_list:
        risques_sorted = sorted(risques_list,
            key=lambda x: int(x.get('criticite', 0)), reverse=True)
        for i, rsk in enumerate(risques_sorted):
            crit = int(rsk.get('criticite', 0))
            if crit >= 12:   niv, fc = 'Critique', 'C0392B'
            elif crit >= 6:  niv, fc = 'Élevé',    'E67E22'
            elif crit >= 3:  niv, fc = 'Modéré',   'F1C40F'
            else:            niv, fc = 'Faible',   '27AE60'
            bg = GRAY4 if i % 2 else WHITE
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            _build_row(t, [
                dict(text=rsk.get('description', ''), fill=bg, size=8),
                dict(text=rsk.get('categorie', ''),   fill=bg, size=8),
                dict(text=str(rsk.get('probabilite', '')), fill=bg, size=8,
                     align=WD_ALIGN_PARAGRAPH.CENTER),
                dict(text=str(rsk.get('impact', '')),      fill=bg, size=8,
                     align=WD_ALIGN_PARAGRAPH.CENTER),
                dict(text=str(crit), fill=fc, bold=True, size=9,
                     color='FFFFFF', align=WD_ALIGN_PARAGRAPH.CENTER),
                dict(text=rsk.get('action', ''), fill=bg, size=8),
                dict(text=rsk.get('statut', ''), fill=PINK, size=8,
                     align=WD_ALIGN_PARAGRAPH.CENTER),
            ])
    else:
        _build_row(t, [dict(text='Aucun risque enregistré dans le registre.',
                            fill=WHITE, size=8, span=7, height=400)])

    if risques_list:
        nb_crit  = sum(1 for r in risques_list if int(r.get('criticite', 0)) >= 12)
        nb_eleve = sum(1 for r in risques_list if 6 <= int(r.get('criticite', 0)) < 12)
        nb_tot   = len(risques_list)
        synthese = (f"Total : {nb_tot} risque(s)   |   "
                    f"Critiques (>=12) : {nb_crit}   |   "
                    f"Élevés (6-11) : {nb_eleve}")
        fill_s = 'C0392B' if nb_crit > 0 else ('E67E22' if nb_eleve > 0 else '27AE60')
        _build_row(t, [dict(text=synthese, fill=fill_s, bold=True,
                            size=8, span=7, color='FFFFFF')])
    br()

    # ── 7. VALIDATION & SIGNATURES ────────────────────────────────────────────
    t = _make_table(doc, CW4)
    _build_row(t, [_sec('7. VALIDATION & SIGNATURES')])
    _build_row(t, [_chdr('Valideur'), _chdr('Nom / Prénom'),
                   _chdr('Date'), _chdr('Visa')])
    for role in ['Chef de projet', 'Responsable DSI', 'Direction']:
        _build_row(t, [_lbl(role), _val('', WHITE), _val('', PINK), _val('', WHITE)],
                   default_height=700)

    # ── Pied de page ──────────────────────────────────────────────────────────
    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.clear()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(
        f"Budget Manager Pro V5 — Fiche projet {p.get('code', '')} — Généré le {today}")
    run.font.name = 'Arial'
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(*_rgb(GRAY2))

    doc.save(output_path)
    logger.info(f"Fiche projet générée : {output_path}")
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# Point d'entrée web (charge les données depuis PostgreSQL)
# ─────────────────────────────────────────────────────────────────────────────
def generer_fiche_depuis_id_pg(projet_id: int, output_dir: str, db) -> str:
    """
    Charge le projet depuis PostgreSQL et génère la fiche Word.
    db         : instance DatabaseService
    output_dir : dossier temporaire de sortie
    Retourne le chemin du fichier .docx créé.
    """
    proj_row = db.fetch_one(
        "SELECT p.*, s.nom as service_nom "
        "FROM projets p LEFT JOIN services s ON s.id=p.service_id "
        "WHERE p.id=%s", [projet_id]
    )
    if not proj_row:
        raise ValueError(f"Projet {projet_id} introuvable")

    proj = dict(proj_row)

    def sg(k, default=''):
        v = proj.get(k)
        return v if v is not None else default

    # ── Tâches ────────────────────────────────────────────────────────────────
    taches_rows = db.fetch_all(
        "SELECT titre, statut, date_echeance, estimation_heures "
        "FROM taches WHERE projet_id=%s ORDER BY date_echeance",
        [projet_id]
    ) or []
    taches = [{
        'titre':    r.get('titre') or '',
        'statut':   r.get('statut') or '',
        'echeance': _fmt_date(r.get('date_echeance')),
        'heures':   f"{int(r.get('estimation_heures') or 0)}h",
    } for r in taches_rows]

    # ── Bons de commande ──────────────────────────────────────────────────────
    bcs_rows = db.fetch_all(
        "SELECT numero_bc, objet, statut, montant_ttc FROM bons_commande "
        "WHERE projet_id=%s ORDER BY date_creation DESC LIMIT 8",
        [projet_id]
    ) or []
    if bcs_rows:
        bcs_str = '\n'.join(
            f"• {r.get('numero_bc', '')} — {(r.get('objet') or '')[:35]} "
            f"({r.get('statut', '')}) — {_fmt_eur(r.get('montant_ttc'))}"
            for r in bcs_rows
        )
    else:
        bcs_str = 'Aucun BC lié'

    # ── Chef de projet & responsable ──────────────────────────────────────────
    def _contact_nom(contact_id):
        if not contact_id:
            return ''
        try:
            r = db.fetch_one(
                "SELECT COALESCE(prenom,'') || ' ' || COALESCE(nom,'') AS n "
                "FROM contacts WHERE id=%s", [contact_id]
            )
            return (r.get('n') or '').strip() if r else ''
        except Exception:
            return ''

    chef      = _contact_nom(sg('chef_projet_contact_id') or None)
    responsable = _contact_nom(sg('responsable_contact_id') or None)

    # ── Équipe (membres) ──────────────────────────────────────────────────────
    equipe = ''
    try:
        rows = db.fetch_all(
            "SELECT COALESCE(pe.membre_label, TRIM(COALESCE(u.prenom,'') || ' ' || COALESCE(u.nom,''))) AS n "
            "FROM projet_equipe pe LEFT JOIN utilisateurs u ON u.id=pe.utilisateur_id "
            "WHERE pe.projet_id=%s", [projet_id]
        ) or []
        equipe = ', '.join(r.get('n') for r in rows if r.get('n'))
    except Exception:
        pass

    # ── Contacts externes ─────────────────────────────────────────────────────
    contacts = []
    try:
        rows = db.fetch_all(
            "SELECT COALESCE(pc.role,'') AS role, "
            "  COALESCE(pc.contact_libre, TRIM(COALESCE(c.prenom,'') || ' ' || COALESCE(c.nom,''))) AS nom, "
            "  COALESCE(c.fonction,'') AS fonction, COALESCE(c.email,'') AS email "
            "FROM projet_contacts pc LEFT JOIN contacts c ON c.id=pc.contact_id "
            "WHERE pc.projet_id=%s", [projet_id]
        ) or []
        contacts = [dict(r) for r in rows]
    except Exception:
        pass

    # ── Prestataires ──────────────────────────────────────────────────────────
    prestataires = ''
    try:
        rows = db.fetch_all(
            "SELECT f.nom FROM projet_prestataires pp "
            "JOIN fournisseurs f ON f.id=pp.fournisseur_id "
            "WHERE pp.projet_id=%s", [projet_id]
        ) or []
        prestataires = ', '.join(r.get('nom', '') for r in rows if r.get('nom'))
    except Exception:
        pass

    data = {
        'code':               sg('code'),
        'nom':                sg('nom'),
        'statut':             sg('statut'),
        'phase':              sg('phase'),
        'priorite':           sg('priorite'),
        'type_projet':        sg('type_projet'),
        'avancement':         sg('avancement', 0),
        'date_debut':         sg('date_debut'),
        'date_fin':           sg('date_fin_prevue'),
        'date_fin_reelle':    sg('date_fin_reelle'),
        'description':        sg('description'),
        'chef_projet':        chef,
        'responsable':        responsable,
        'equipe':             equipe,
        'prestataires':       prestataires,
        'service':            sg('service_nom'),
        'budget_previsionnel': proj.get('budget_initial') or proj.get('budget_estime'),
        'budget_vote':        proj.get('budget_actuel'),
        'budget_consomme':    proj.get('budget_consomme'),
        'ligne_budg':         '',
        'bcs':                bcs_str,
        'objectifs':          sg('objectifs'),
        'enjeux':             sg('enjeux'),
        'risques':            sg('risques'),
        'gains':              sg('gains'),
        'contraintes':        sg('contraintes'),
        'solutions':          sg('solutions'),
        'financement':        sg('financement'),
        'taches':             taches,
        'contacts':           contacts,
        'registre_risques':   sg('registre_risques'),
        'contraintes_6axes':  sg('contraintes_6axes'),
        'triangle_tensions':  sg('triangle_tensions'),
        'arbitrage':          sg('arbitrage'),
    }

    code = sg('code', f'PRJ{projet_id}')
    out_path = os.path.join(output_dir, f"fiche_projet_{code}.docx")
    return generer_fiche_projet(data, out_path)
